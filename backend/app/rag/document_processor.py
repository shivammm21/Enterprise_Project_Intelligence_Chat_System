import base64
import logging
from pathlib import Path
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}


# ── Text extractors ────────────────────────────────────────────────────────────

def extract_text_from_pdf(filepath: str) -> List[Document]:
    from pypdf import PdfReader
    reader = PdfReader(filepath)
    docs = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": Path(filepath).name, "page": page_num + 1},
            ))
    return docs


def extract_text_from_docx(filepath: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(
        page_content=full_text,
        metadata={"source": Path(filepath).name, "page": 1},
    )]


def extract_text_from_txt(filepath: str) -> List[Document]:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={"source": Path(filepath).name, "page": 1},
    )]


# ── Image extractors ───────────────────────────────────────────────────────────

_IMAGE_PROMPT = (
    "You are analyzing a project-related image (e.g., architecture diagram, system design, "
    "flowchart, ERD, UI mockup, screenshot, or any visual document). "
    "Provide a detailed description covering: all components/services/elements shown, "
    "their relationships and data flows, labels and annotations, technology stack references, "
    "and any key design decisions visible. "
    "Be thorough — this description will be used to answer questions about the project."
)


def _describe_image_with_openai(filepath: str, api_key: str) -> str:
    from openai import OpenAI

    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
        "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp",
        "tiff": "image/tiff", "tif": "image/tiff",
    }
    ext = Path(filepath).suffix.lower().lstrip(".")
    mime_type = mime_map.get(ext, "image/png")

    with open(filepath, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": _IMAGE_PROMPT},
                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}},
            ],
        }],
        max_tokens=1500,
    )
    return response.choices[0].message.content or ""


def _describe_image_with_gemini(filepath: str, api_key: str) -> str:
    import requests

    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
        "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp",
        "tiff": "image/tiff", "tif": "image/tiff",
    }
    ext = Path(filepath).suffix.lower().lstrip(".")
    mime_type = mime_map.get(ext, "image/png")

    with open(filepath, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
    body = {
        "contents": [{
            "parts": [
                {"text": _IMAGE_PROMPT},
                {"inline_data": {"mime_type": mime_type, "data": image_data}},
            ]
        }]
    }
    resp = requests.post(url, json=body, params={"key": api_key}, timeout=60)
    resp.raise_for_status()
    return resp.json()["candidates"][0]["content"]["parts"][0]["text"]


def extract_text_from_image(filepath: str) -> List[Document]:
    """
    Describe an image using the active vision LLM (OpenAI GPT-4o or Gemini 1.5 Flash).
    Provider is chosen automatically based on which API key is configured.
    Raises on failure so the caller gets a real error message.
    """
    from app.config import settings

    filename = Path(filepath).name
    description = ""

    if settings.AI_PROVIDER == "openai":
        logger.info(f"Describing image via OpenAI GPT-4o: {filename}")
        description = _describe_image_with_openai(filepath, settings.OPENAI_API_KEY)
    else:
        logger.info(f"Describing image via Gemini 1.5 Flash: {filename}")
        description = _describe_image_with_gemini(filepath, settings.GOOGLE_API_KEY)

    if not description.strip():
        raise ValueError(f"Vision API returned empty description for {filename}")

    return [Document(
        page_content=description,
        metadata={
            "source": filename,
            "page": 1,
            "content_type": "image",
        },
    )]


# ── Dispatcher ─────────────────────────────────────────────────────────────────

def load_document(filepath: str) -> List[Document]:
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    elif ext in IMAGE_EXTENSIONS:
        return extract_text_from_image(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def split_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)
